import os
import json
import re
import time
import tempfile
import subprocess
import traceback
from flask import Flask, request, jsonify, render_template
import vertexai
from vertexai.generative_models import GenerativeModel, GenerationConfig
from pypdf import PdfReader
from docx import Document

app = Flask(__name__)

# --- CONFIGURATION ---
PROJECT_ID = "project-01-test-488412" 
REGION = "us-central1"

# Initialize Vertex AI with explicit credentials fallback
# Cloud Shell's metadata server sometimes fails to provide service account email
import google.auth
try:
    credentials, project = google.auth.default()
    vertexai.init(project=PROJECT_ID, location=REGION, credentials=credentials)
    print(f"[AUTH] Vertex AI initialized with credentials for project: {PROJECT_ID}")
except Exception as e:
    print(f"[AUTH] Warning: google.auth.default() failed ({e}), trying basic init...")
    vertexai.init(project=PROJECT_ID, location=REGION)

# --- MODEL CASCADE ---
# gemini-2.5-flash: Fast, reliable, high quota — ideal for batch CV screening
# gemini-2.5-pro: Fallback if flash is unavailable
# gemini-2.0-flash: Emergency fallback — stable, high quota
MODELS = [
    GenerativeModel("gemini-2.5-flash"),
    GenerativeModel("gemini-2.5-pro"),
    GenerativeModel("gemini-2.0-flash"),
]

MAX_CV_CHARS = 8000
MAX_JD_CHARS = 5000


def extract_text_from_file(file):
    filename = file.filename.lower()
    try:
        file.seek(0)
        
        if filename.endswith('.pdf'):
            reader = PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            print(f"[EXTRACT] PDF '{filename}': {len(text)} chars")
            return text

        elif filename.endswith('.doc') and not filename.endswith('.docx'):
            # Legacy .doc format — try multiple extraction strategies
            print(f"[EXTRACT] DOC '{filename}': Legacy .doc detected, trying extraction...")
            file_bytes = file.read()
            
            with tempfile.TemporaryDirectory() as tmpdir:
                doc_path = os.path.join(tmpdir, "input.doc")
                with open(doc_path, 'wb') as f:
                    f.write(file_bytes)
                
                # Strategy 1: LibreOffice (best quality, may not be installed)
                text = _try_libreoffice(doc_path, tmpdir, filename)
                if text:
                    return text
                
                # Strategy 2: antiword (lightweight, may not be installed)
                text = _try_antiword(doc_path, filename)
                if text:
                    return text
                
                # Strategy 3: Pure Python binary extraction (always works, lower quality)
                text = _try_binary_doc_extract(file_bytes, filename)
                if text:
                    return text
                
                print(f"[ERROR] All .doc extraction strategies failed for '{filename}'")
                return ""

        elif filename.endswith('.docx'):
            return _extract_from_docx(file, filename)

        else:
            file.seek(0) 
            return file.read().decode('utf-8', errors='ignore')
    except Exception as e:
        print(f"Extraction error on {filename}: {e}")
        return ""


def _extract_from_docx(source, filename):
    """Extract text from a .docx file (accepts file object or file path)"""
    try:
        doc = Document(source)

        para_text = "\n".join([p.text for p in doc.paragraphs]).strip()

        table_lines = []
        seen = set()
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    line = " | ".join(row_cells)
                    if line not in seen:
                        seen.add(line)
                        table_lines.append(line)
        table_text = "\n".join(table_lines)

        print(f"[EXTRACT] DOCX '{filename}': paras={len(para_text)} chars, tables={len(table_text)} chars")

        if len(para_text) > 100:
            extra_table = []
            for line in table_lines:
                snippet = line[:50].lower()
                if snippet not in para_text.lower():
                    extra_table.append(line)
            if extra_table:
                return para_text + "\n\n" + "\n".join(extra_table)
            return para_text
        else:
            result = table_text if table_text else para_text
            if not result.strip():
                print(f"[WARNING] DOCX '{filename}': EMPTY extraction!")
            return result
    except Exception as e:
        print(f"DOCX extraction error on {filename}: {e}")
        return ""


def _try_libreoffice(doc_path, tmpdir, filename):
    """Strategy 1: Convert .doc to .docx via LibreOffice"""
    try:
        result = subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'docx', doc_path, '--outdir', tmpdir],
            capture_output=True, text=True, timeout=30
        )
        docx_path = os.path.join(tmpdir, "input.docx")
        if os.path.exists(docx_path):
            text = _extract_from_docx(docx_path, filename)
            if text and len(text.strip()) > 50:
                print(f"[EXTRACT] DOC '{filename}': LibreOffice strategy succeeded ({len(text)} chars)")
                return text
    except FileNotFoundError:
        print(f"[EXTRACT] DOC '{filename}': LibreOffice not installed, skipping...")
    except Exception as e:
        print(f"[EXTRACT] DOC '{filename}': LibreOffice failed: {e}")
    return ""


def _try_antiword(doc_path, filename):
    """Strategy 2: Extract text via antiword (lightweight CLI tool)"""
    try:
        result = subprocess.run(
            ['antiword', doc_path],
            capture_output=True, text=True, timeout=15
        )
        if result.returncode == 0 and len(result.stdout.strip()) > 50:
            print(f"[EXTRACT] DOC '{filename}': antiword strategy succeeded ({len(result.stdout)} chars)")
            return result.stdout.strip()
    except FileNotFoundError:
        print(f"[EXTRACT] DOC '{filename}': antiword not installed, skipping...")
    except Exception as e:
        print(f"[EXTRACT] DOC '{filename}': antiword failed: {e}")
    return ""


def _try_binary_doc_extract(file_bytes, filename):
    """Strategy 3: Pure Python — extract UTF-16LE text runs from binary .doc file.
    No external dependencies. Lower quality but always works."""
    try:
        text_runs = []
        current_run = []
        
        for i in range(0, len(file_bytes) - 1, 2):
            lo, hi = file_bytes[i], file_bytes[i + 1]
            if hi == 0 and (32 <= lo <= 126 or lo in (9, 10, 13)):
                current_run.append(chr(lo))
            else:
                if len(current_run) > 20:
                    text_runs.append(''.join(current_run).strip())
                current_run = []
        
        if len(current_run) > 20:
            text_runs.append(''.join(current_run).strip())
        
        # Sort by length — the actual document text is usually the longest runs
        text_runs.sort(key=len, reverse=True)
        
        # Take top runs that look like real content (not binary metadata)
        clean_runs = []
        for run in text_runs[:15]:
            # Skip runs that look like metadata/binary noise
            alpha_ratio = sum(1 for c in run if c.isalpha()) / max(len(run), 1)
            if alpha_ratio > 0.4 and len(run) > 30:
                clean_runs.append(run)
        
        result = '\n'.join(clean_runs)
        
        if len(result.strip()) > 50:
            print(f"[EXTRACT] DOC '{filename}': Binary extraction succeeded ({len(result)} chars)")
            return result
    except Exception as e:
        print(f"[EXTRACT] DOC '{filename}': Binary extraction failed: {e}")
    return ""


def clean_and_parse_json(raw_response):
    try:
        match = re.search(r'\{.*\}', raw_response, re.DOTALL)
        if match:
            raw_response = match.group(0)
        
        cleaned_str = raw_response.replace('\n', ' ').replace('\r', '').strip()
        data = json.loads(cleaned_str)
        
        if isinstance(data, list):
            data = data[0]
            
        return data
    except Exception as e:
        print(f"JSON Parsing Error: {e}")
        raise ValueError(f"Failed to parse AI response: {str(e)}")


def call_gemini_with_cascade(prompt, max_retries_per_model=2):
    """
    Try each model in the cascade. For each model, retry with backoff.
    This ensures if gemini-2.5-flash is down, we fall to 2.5-pro, then 1.5-flash.
    """
    last_error = None
    
    for model_idx, model in enumerate(MODELS):
        model_name = model._model_name if hasattr(model, '_model_name') else f"model-{model_idx}"
        
        for attempt in range(max_retries_per_model):
            try:
                print(f"[GEMINI] Trying {model_name} (attempt {attempt+1}/{max_retries_per_model})...")
                response = model.generate_content(prompt)
                print(f"[GEMINI] Success with {model_name}")
                return response
                
            except Exception as e:
                last_error = e
                error_str = str(e).lower()
                is_retryable = any(kw in error_str for kw in [
                    "503", "429", "resource exhausted", "overloaded",
                    "rate limit", "quota", "unavailable", "service_unavailable",
                    "internal", "deadline"
                ])
                
                if is_retryable:
                    if attempt < max_retries_per_model - 1:
                        wait_time = (attempt + 1) * 3  # 3s, 6s
                        print(f"[GEMINI] {model_name} returned {str(e)[:80]}. Retrying in {wait_time}s...")
                        time.sleep(wait_time)
                    else:
                        print(f"[GEMINI] {model_name} exhausted retries. Moving to next model...")
                        break  # Move to next model in cascade
                else:
                    # Non-retryable error (bad prompt, auth, etc.) — don't retry
                    print(f"[GEMINI] Non-retryable error on {model_name}: {str(e)[:100]}")
                    raise e
    
    # All models exhausted
    raise Exception(f"All Gemini models failed. Last error: {str(last_error)}")


@app.route('/')
def index():
    return render_template('index.html')


@app.route('/api/screen', methods=['POST'])
def screen_cv():
    try:
        jd_file = request.files.get('jd')
        cv_file = request.files.get('cv')
        recruiter_notes = request.form.get('notes', '')

        if not cv_file:
            return jsonify({"success": False, "error": "CV file missing"}), 400

        jd_text = extract_text_from_file(jd_file) if jd_file else "Not provided"
        cv_text = extract_text_from_file(cv_file)

        # Check for blank/empty document
        if not cv_text or len(cv_text.strip()) == 0:
            return jsonify({
                "success": False, 
                "error": f"'{cv_file.filename}' appears to be a blank document with no content."
            }), 400

        # Check for image-only document (very little extractable text = likely scanned/image PDF)
        if len(cv_text.strip()) < 50:
            return jsonify({
                "success": False, 
                "error": f"'{cv_file.filename}' appears to contain only images. Please upload a text-based document."
            }), 400

        # Same check for JD
        if jd_file and (not jd_text or len(jd_text.strip()) == 0):
            return jsonify({
                "success": False, 
                "error": f"'{jd_file.filename}' appears to be a blank document with no content."
            }), 400

        if jd_file and len(jd_text.strip()) < 50:
            return jsonify({
                "success": False, 
                "error": f"'{jd_file.filename}' appears to contain only images. Please upload a text-based JD."
            }), 400

        jd_text = jd_text[:MAX_JD_CHARS]
        cv_text = cv_text[:MAX_CV_CHARS]

        screening_prompt = f"""
            Act as a Strategic Talent Architect. 
            TASK: Conduct a forensic audit of the CV against the JD.

            CORE INSTRUCTION: 
            1. Identify the Candidate's actual name from the CV content (usually found at the top). 
            2. If no name is clearly identifiable, use the filename provided in context.

            RECRUITER OVERRIDES:
            {recruiter_notes if recruiter_notes else "None."}

            CRITICAL FORMATTING RULES:
            - Every string in every array MUST be a single concise sentence, maximum 15 words. No long explanations.
            - Each array MUST have at most 3 items. Prioritize the most important points only.
            - No markdown formatting whatsoever. No bold (**), no headers (##), no bullet symbols (-, *).
            - Write in plain, crisp, professional English only.
            - The rationale must be one single sentence, maximum 20 words.

            OUTPUT ONLY VALID JSON:
            {{
                "candidate_name": "Extract real name from CV here",
                "overallScore": 0-100,
                "recommendation": "Hierarchy Level",
                "rationale": "One crisp sentence, max 20 words.",
                "strengths": {{
                    "NIRF_and_Pedigree": ["max 2 items, each under 15 words"],
                    "Experience_Alignment": ["max 2 items, each under 15 words"],
                    "Projects_and_Quantifiable_Impact": ["max 2 items, each under 15 words"]
                }},
                "proximity_matches": ["max 2 items, each under 15 words"],
                "gaps": {{ "Functional_Gaps": ["max 2 items, each under 15 words"], "Domain_Mismatch": ["max 1 item, under 15 words"] }},
                "jd_enhancement": {{ "missing_in_jd": [] }}
            }}

            JD: {jd_text}
            CV FILENAME: {cv_file.filename}
            CV: {cv_text}
        """

        response = call_gemini_with_cascade(screening_prompt)
        parsed_data = clean_and_parse_json(response.text)

        return jsonify({"success": True, "data": parsed_data})

    except Exception as e:
        print(f"Error: {traceback.format_exc()}")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/enhance-jd', methods=['POST'])
def enhance_jd():
    try:
        jd_file = request.files.get('jd_file')
        recruiter_notes = request.form.get('notes', '')

        if not jd_file:
            return jsonify({"success": False, "error": "No file uploaded"}), 400

        base_jd_text = extract_text_from_file(jd_file)
        base_jd_text = base_jd_text[:MAX_JD_CHARS]

        enhancer_prompt = f"""
        Rewrite this JD into a high-performance Job Description.
        Notes: {recruiter_notes}

        CRITICAL FORMATTING RULE: The output must be plain text only.
        Do NOT use any markdown syntax — no ## headers, no ** bold **, no * bullets, no --- dividers.
        Use simple line breaks and paragraph spacing for structure.
        Use UPPERCASE for section headings instead of markdown.

        Return JSON format: {{"enhanced_text": "Plain text content here, absolutely no markdown formatting"}}
        
        BASE JD: {base_jd_text}
        """

        response = call_gemini_with_cascade(enhancer_prompt)
        parsed_result = clean_and_parse_json(response.text)

        return jsonify({"success": True, "enhanced_text": parsed_result.get('enhanced_text', '')})

    except Exception as e:
        print(f"Enhancement Error: {str(e)}")
        return jsonify({"success": False, "error": str(e)}), 500

        
if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=int(os.environ.get("PORT", 8080)))